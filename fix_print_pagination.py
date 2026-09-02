from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX = Path("/Users/EvanBelfi/Documents/Projects/USGA-GC-Dashboard/Joint Strategy and Project Accelerate Alignment1 - Print Fixed.docx")


doc = Document(DOCX)

# Prevent Word Online from stranding the Section 3 heading at the bottom of a
# page. Anchor the heading, lead sentence, and subheading to the table that
# follows, without imposing a hard page break.
section3 = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Areas of Overlap and Potential Work Team Consolidations")
potential = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Potential Consolidations:")
for i in range(section3, potential + 1):
    doc.paragraphs[i].paragraph_format.keep_with_next = True
    doc.paragraphs[i].paragraph_format.keep_together = True

# Let Section 4 flow naturally in the digital document, while compacting it
# just enough to remain on the same printed page across renderers.
start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "What Greater Alignment Could Bring")
end = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Stronger connection between strategy and execution")
for i in range(start, end + 1):
    p = doc.paragraphs[i]
    p.paragraph_format.keep_with_next = i == start
    p.paragraph_format.keep_together = True
    if i > start:
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    if 18 <= i <= end:
        p.paragraph_format.space_before = Pt(0)

# Remove the forced page start so the digital reading layout remains natural.
section5 = next(p for p in doc.paragraphs if p.text.strip() == "From Separate Reporting to Shared Membership Intelligence")
section5.paragraph_format.page_break_before = False
section5.paragraph_format.keep_with_next = True

# Remove the empty list paragraph between Sections 4 and 5; it can acquire
# printer-dependent list spacing even though it looks blank on screen.
for p in list(doc.paragraphs):
    if p._p.getprevious() is not None and not p.text.strip() and p.style.name == "List Paragraph":
        prev_text = p._p.getprevious().xpath("string(.)").strip()
        next_el = p._p.getnext()
        next_text = next_el.xpath("string(.)").strip() if next_el is not None else ""
        if "Stronger connection between strategy and execution" in prev_text and "From Separate Reporting" in next_text:
            p._element.getparent().remove(p._element)
            break

doc.save(DOCX)
